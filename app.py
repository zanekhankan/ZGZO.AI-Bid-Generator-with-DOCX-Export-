import streamlit as st
import json
import os
from docx import Document
from docx.shared import Pt
from datetime import datetime
import pandas as pd

st.set_page_config(page_title="ZGZO.AI Bid Generator", layout="centered")
st.title("📄 ZGZO.AI - AI Bid Generator")

# -------------------------
# GC PROFILE SELECTION
# -------------------------
gc_dir = "gc_profiles"
os.makedirs(gc_dir, exist_ok=True)

st.subheader("1. Select GC Profile")
gc_files = [f for f in os.listdir(gc_dir) if f.endswith("_config.json")]

if not gc_files:
    st.warning("No GC profiles found. Please create one in the Config Creator first.")
else:
    selected_gc = st.selectbox("Choose GC Profile", gc_files)

    # -------------------------
    # FILE UPLOAD
    # -------------------------
    st.subheader("2. Upload Specs or Drawings")
    uploaded_file = st.file_uploader("Upload PDF or DOCX file", type=["pdf", "docx"])

    # -------------------------
    # PRICING METHOD TOGGLE
    # -------------------------
    use_manual = st.radio("Select Pricing Method", ["Use Markup", "Enter Prices Manually"])

    # -------------------------
    # MARKUP CONTROL
    # -------------------------
    global_markup = st.number_input("Global Markup % (only used in markup mode)", min_value=0.0, value=10.0)

    # TEMP FAKE DATA (replace with real extracted items later)
    line_items = [
        {"Description": "Concrete Slab", "Quantity": 100, "Unit": "sqft"},
        {"Description": "Rebar", "Quantity": 50, "Unit": "lbs"},
        {"Description": "Excavation", "Quantity": 200, "Unit": "cubic ft"},
    ]

    manual_prices = []

    if use_manual == "Enter Prices Manually":
        st.markdown("### Manual Price Entry")

        for idx, item in enumerate(line_items):
            st.write(f"**{item['Description']}**")

            qty = item.get("Quantity", 1)
            unit = item.get("Unit", "")
            unit_price = st.number_input(
                f"Unit Price for {item['Description']}",
                min_value=0.0,
                value=0.0,
                key=f"price_{idx}"
            )
            total = qty * unit_price
            st.write(f"Quantity: {qty} {unit}, Total: ${total:,.2f}")

            manual_prices.append({
                "Description": item["Description"],
                "Quantity": qty,
                "Unit": unit,
                "Unit Price": unit_price,
                "Total": total
            })

        subtotal = sum(row["Total"] for row in manual_prices)
        tax = st.number_input("Tax %", min_value=0.0, max_value=100.0, value=8.0)
        total_with_tax = subtotal * (1 + tax / 100)

        st.markdown(f"### Subtotal: ${subtotal:,.2f}")
        st.markdown(f"### Total with Tax: ${total_with_tax:,.2f}")

        # Export to Excel
        df = pd.DataFrame(manual_prices)
        df["Total"] = df["Total"].round(2)
        df["Unit Price"] = df["Unit Price"].round(2)

        excel_file = df.to_csv(index=False).encode('utf-8')
        st.download_button("📥 Download Excel", data=excel_file, file_name="ZGZO_Estimate.csv", mime="text/csv")

        # Save bid
        if st.button("💾 Save This Bid"):
            os.makedirs("saved_bids", exist_ok=True)
            bid_name = f"Bid_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
            with open(os.path.join("saved_bids", bid_name), "w") as f:
                json.dump(manual_prices, f, indent=2)
            st.success(f"Bid saved as {bid_name}")

        # Load bid
        st.subheader("📂 Load Previous Bid")
        saved_bids = os.listdir("saved_bids") if os.path.exists("saved_bids") else []
        if saved_bids:
            selected_bid = st.selectbox("Select a saved bid", saved_bids)
            if st.button("Load Selected Bid"):
                with open(os.path.join("saved_bids", selected_bid), "r") as f:
                    manual_prices = json.load(f)
                st.success(f"Loaded {selected_bid}")
        else:
            st.info("No saved bids available.")

    # -------------------------
    # GENERATE BID DOCUMENT
    # -------------------------
    if uploaded_file and selected_gc:
        with open(os.path.join(gc_dir, selected_gc), "r") as f:
            config = json.load(f)

        st.success(f"Using profile: {config['gc_name']}")
        st.write(f"License #: {config['license']}")
        st.write(f"Markup: {config['markup_percent']}%")
        st.write(f"Tone: {config['tone'].capitalize()}")

        if st.button("Generate Bid Document"):
            doc = Document()
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Arial'
            font.size = Pt(11)

            doc.add_heading("Bid Proposal", 0)
            doc.add_paragraph("Project: Test Project")
            doc.add_paragraph("Prepared for: City of Testville")
            doc.add_paragraph(f"Prepared by: {config['gc_name']}")
            doc.add_paragraph(f"License #: {config['license']}")
            doc.add_paragraph(f"Contact: {config['contact']} | {config['phone']}")
            doc.add_paragraph("Date: " + datetime.today().strftime('%B %d, %Y'))
            doc.add_page_break()

            doc.add_heading("Scope of Work", level=1)
            scope_items = [item["Description"] for item in manual_prices] if use_manual == "Enter Prices Manually" else [
                "Division 02: Selective demolition of tile and plumbing fixtures",
                "Division 03: New slab pour for restroom flooring",
                "Division 09: New ceramic wall tile and paint finishes",
                "Division 15: Installation of ADA-compliant fixtures and hot water piping",
                "Division 16: Relocate lighting, add occupancy sensors"
            ]
            for item in scope_items:
                doc.add_paragraph(f"- {item}")

            doc.add_heading("Cost Estimate Summary", level=1)
            table = doc.add_table(rows=1, cols=3)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Division'
            hdr_cells[1].text = 'Description'
            hdr_cells[2].text = 'Estimated Cost'

            if use_manual == "Enter Prices Manually":
                for idx, item in enumerate(manual_prices):
                    row_cells = table.add_row().cells
                    row_cells[0].text = f"{idx+1:02}"
                    row_cells[1].text = item["Description"]
                    row_cells[2].text = f"${item['Total']:,.2f}"
            else:
                default_cost_data = [
                    ("02", "Demolition", "$5,000"),
                    ("03", "Concrete", "$12,000"),
                    ("09", "Finishes", "$8,000"),
                    ("15", "Plumbing", "$10,000"),
                    ("16", "Electrical", "$7,000")
                ]
                for div, desc, cost in default_cost_data:
                    row_cells = table.add_row().cells
                    row_cells[0].text = div
                    row_cells[1].text = desc
                    row_cells[2].text = cost

            doc.add_heading("Project Timeline", level=1)
            doc.add_paragraph("Phase 1 – Demolition: Week 1")
            doc.add_paragraph("Phase 2 – Concrete & Plumbing: Weeks 2–3")
            doc.add_paragraph("Phase 3 – Electrical & Finishes: Weeks 4–5")

            doc.add_heading("Inclusions & Exclusions", level=1)
            doc.add_paragraph("Inclusions:")
            doc.add_paragraph("- Labor, material, and equipment to complete scope", style='List Bullet')
            doc.add_paragraph("- Cleanup and disposal")
            doc.add_paragraph("Exclusions:")
            doc.add_paragraph("- Permits and inspection fees", style='List Bullet')
            doc.add_paragraph("- Hazardous material abatement")

            doc.add_heading("Acknowledgement", level=1)
            doc.add_paragraph(config["gc_name"])
            doc.add_paragraph("Authorized Estimator")
            doc.add_paragraph("Signature: _________________________")

            doc.add_heading("Legal Notice", level=1)
            doc.add_paragraph(config["legal"])

            output_path = "ZGZO_AI_Bid_Output.docx"
            doc.save(output_path)

            with open(output_path, "rb") as file:
                st.download_button(
                    label="Download Bid Document",
                    data=file,
                    file_name="ZGZO_AI_Bid_Output.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
